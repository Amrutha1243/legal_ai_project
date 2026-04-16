import os
import random
from textwrap import fill

import docx # Already used in vectorstore_manager

try:
    from fpdf import FPDF
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "fpdf2"])
    from fpdf import FPDF

DATA_DIR = os.path.dirname(os.path.abspath(__file__))

# 20 Distinct Acts for the Knowledge Base
ACTS = [
    ("Special_Marriage_Act_1954", "txt"),
    ("Companies_Act_2013", "pdf"),
    ("Motor_Vehicles_Amendment_Act_2019", "docx"),
    ("Information_Technology_Rules_2021", "txt"),
    ("Prevention_of_Money_Laundering_Act_2002", "pdf"),
    ("Right_to_Information_Act_2005", "docx"),
    ("Hindu_Marriage_Act_1955", "txt"),
    ("Indian_Contract_Act_1872", "pdf"),
    ("Arbitration_and_Conciliation_Act_1996", "docx"),
    ("Specific_Relief_Act_1963", "txt"),
    ("Juvenile_Justice_Act_2015", "pdf"),
    ("Prevention_of_Corruption_Act_1988", "docx"),
    ("Competition_Act_2002", "txt"),
    ("Foreign_Exchange_Management_Act_1999", "pdf"),
    ("Insolvency_and_Bankruptcy_Code_2016", "docx"),
    ("Securitisation_and_Reconstruction_of_Financial_Assets_Act_2002", "txt"),
    ("Real_Estate_Regulation_and_Development_Act_2016", "pdf"),
    ("Protection_of_Children_from_Sexual_Offences_Act_2012", "docx"),
    ("Digital_Personal_Data_Protection_Act_2023", "txt"),
    ("Labour_Codes_and_Occupational_Safety_2020", "pdf"),
]

# Base template text to expand upon to easily span 1000+ lines
PARAGRAPH_TEMPLATES = [
    "This section outlines the primary regulatory mechanisms applied to citizens and entities operating under the jurisdiction of the Republic of India. Any deviation from the prescribed norms shall result in immediate notification to the relevant central authority, initiating an investigation spanning not more than 90 days. During this period, the accused entity must surrender all relevant documentation.",
    "The designated tribunal holds the absolute authority to request financial records, digital communications, and server logs. If an individual refuses compliance under subsection {i}, they will be subject to a fine not exceeding Rs. {fine} and possible imprisonment for a term up to {years} years. Courts are advised to expedite these matters recognizing the public interest.",
    "No civil court shall have jurisdiction to entertain any suit or proceeding in respect of any matter which the adjudicating authority is empowered by or under this Act to determine. Furthermore, no injunction shall be granted by any court or other authority in respect of any action taken or to be taken in pursuance of any power conferred by or under this Act.",
    "Where an offense under this Act has been committed by a company, every person who, at the time the offense was committed, was in charge of, and was responsible to, the company for the conduct of the business of the company, as well as the company itself, shall be deemed to be guilty of the offense and shall be liable to be proceeded against and punished accordingly.",
    "It is crucial to note that the central government retains the power, by notification in the Official Gazette, to amend, alter, or remove constraints stipulated under this specific clause, provided that a draft of such a notification is laid before both Houses of Parliament for a statutory period."
]

def generate_text_content(title):
    lines = [f"THE {title.replace('_', ' ').upper()}"]
    lines.append("An exhaustive codification covering fundamental definitions, regulations, liabilities, and amendments.")
    lines.append("="*80)
    lines.append("\nPRELIMINARY\n")
    
    # Generate ~200 sections to easily cross 1000 lines
    for i in range(1, 251):
        lines.append(f"\nSECTION {i}: Regulations, Procedures and Penalties")
        lines.append("-" * 40)
        
        paragraphs = []
        for j in range(3):
            tmpl = random.choice(PARAGRAPH_TEMPLATES)
            para = tmpl.format(i=i, fine=random.randint(10, 500)*1000, years=random.randint(1, 10))
            # Wrap text lightly to create more lines
            paragraphs.append(fill(para, width=80))
            
        lines.append("\n\n".join(paragraphs))
        lines.append(f"Explanation to Section {i}: For the purposes of this segment, 'compliance' implies absolute adherence without caveats.")
        lines.append(f"Sub-clause ({i}.a): The burden of proof remains intrinsically tied to the primary respondent.")
        
    return "\n".join(lines)

for act_name, ext in ACTS:
    filepath = os.path.join(DATA_DIR, f"{act_name}.{ext}")
    content = generate_text_content(act_name)
    
    print(f"Generating {filepath} ({len(content.splitlines())} lines)...")
    
    if ext == "txt":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
            
    elif ext == "docx":
        doc = docx.Document()
        doc.add_heading(act_name.replace("_", " "), 0)
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        doc.save(filepath)
        
    elif ext == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        # Using multi_cell to handle wrapping and utf-8 encoding natively handled usually by standard ascii
        safe_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=safe_content)
        pdf.output(filepath)

print("All 20 files successfully generated in app/data!")
