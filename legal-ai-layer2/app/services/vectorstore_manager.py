import os
import hashlib
import pickle

if "GOOGLE_API_KEY" not in os.environ and "GEMINI_API_KEY" in os.environ:
    os.environ["GOOGLE_API_KEY"] = os.environ["GEMINI_API_KEY"]

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import GoogleGenerativeAIEmbeddings
from langchain.vectorstores import FAISS

import PyPDF2
import docx

DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")
FAISS_INDEX_PATH = os.path.join(DATA_DIR, "faiss_index")
HASH_FILE = os.path.join(DATA_DIR, "file_hashes.pkl")

def get_file_hash(filepath):
    hasher = hashlib.md5()
    with open(filepath, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
    return text

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
    return text

def load_documents_from_directory():
    documents = []
    
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
        
    for filename in os.listdir(DATA_DIR):
        filepath = os.path.join(DATA_DIR, filename)
        if not os.path.isfile(filepath):
            continue
            
        text = ""
        if filename.endswith(".txt"):
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(filepath)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(filepath)
            
        if text.strip():
            documents.append(Document(page_content=text, metadata={"source": filename}))
            
    return documents

def build_or_update_vectorstore():
    # Load old hashes
    old_hashes = {}
    if os.path.exists(HASH_FILE):
        with open(HASH_FILE, 'rb') as f:
            old_hashes = pickle.load(f)
            
    # Calculate new hashes
    new_hashes = {}
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
        
    for filename in os.listdir(DATA_DIR):
        filepath = os.path.join(DATA_DIR, filename)
        if os.path.isfile(filepath):
            if filename.endswith(".txt") or filename.endswith(".pdf") or filename.endswith(".docx"):
                new_hashes[filename] = get_file_hash(filepath)
            
    # Check if anything changed
    if old_hashes == new_hashes and os.path.exists(FAISS_INDEX_PATH):
        print("Knowledge base is up to date. Loading from disk...")
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        return FAISS.load_local(FAISS_INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
        
    print("Changes detected in data directory. Rebuilding knowledge base...")
    documents = load_documents_from_directory()
    
    if not documents:
        print("No valid documents found. returning None.")
        return None
        
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    split_docs = splitter.split_documents(documents)
    
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vectorstore = FAISS.from_documents(split_docs, embeddings)
    
    # Save to disk
    vectorstore.save_local(FAISS_INDEX_PATH)
    with open(HASH_FILE, 'wb') as f:
        pickle.dump(new_hashes, f)
        
    print("Knowledge base updated successfully.")
    return vectorstore

def get_persistent_vectorstore():
    if os.path.exists(FAISS_INDEX_PATH):
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        try:
            return FAISS.load_local(FAISS_INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
        except ValueError:
            # Langchain 0.0.x vs Langchain 0.1.x backwards compatibility for load_local
            return FAISS.load_local(FAISS_INDEX_PATH, embeddings)
    return build_or_update_vectorstore()

# Auto-run on import to ensure we initialize or update if needed
global_vectorstore = build_or_update_vectorstore()
